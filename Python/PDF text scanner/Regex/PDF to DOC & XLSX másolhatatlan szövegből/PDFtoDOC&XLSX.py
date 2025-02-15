import PyPDF2
import os
import re
from docx import Document
from datetime import datetime
import pytesseract
from pdf2image import convert_from_path
import openpyxl


folder_name = "ready"
current_directory = os.getcwd()

full_path = os.path.join(current_directory, folder_name)
now = datetime.now().date()


if os.path.exists(full_path) and os.path.isdir(full_path):
    print("Folder already exists")
else:
    os.mkdir(full_path)
    print("Folder created")


def replace_text_with_variables(document, replace_dict):
    for paragraph in document.paragraphs:
        for key, val in replace_dict.items():
            if key in paragraph.text:
                paragraph.text = paragraph.text.replace(key, val)
    return document


def extract_text_and_images_from_pdf(file_path):
    images = convert_from_path(file_path)
    text = ""
    for image in images:
        extracted_text = pytesseract.image_to_string(image)
        if extracted_text:
            text += extracted_text + "\n"
    with open(file_path, 'rb') as pdf_file_obj:
        pdf_reader = PyPDF2.PdfReader(pdf_file_obj)
        for page_num in range(len(pdf_reader.pages)):
            page_obj = pdf_reader.pages[page_num]
            if "/XObject" in page_obj["/Resources"]:
                xobject = page_obj["/Resources"]["/XObject"]
                for obj in xobject:
                    if xobject[obj]["/Subtype"] == "/Image":
                        image = xobject[obj].get_object()
                        image_data = image._data
                        # Process the image data as needed
                        # You can save the image or perform further processing here
    return text



def list_pdf_files(directory):
    return [file for file in os.listdir(directory) if file.endswith('.pdf')]


pdf_files = list_pdf_files(current_directory)

for pdf_file in pdf_files:
    full_pdf_path = os.path.join(current_directory, pdf_file)
    text = extract_text_and_images_from_pdf(full_pdf_path)


    invoiceRegex = re.compile("[A-Z]\d[A-z]\d\d\d\d\d")
    inv = re.search(invoiceRegex, text)
    szamla = inv.group() if inv else "No match"




    BLRegex = re.compile("[A-Z][A-Z][A-Z][A-Z][A-Z]\d\d\d\d\d\d\d")
    blmatch = re.search(BLRegex, text)
    bl = blmatch.group() if blmatch else "No match"

    CLLRegex = re.compile(r"\d+ case\(s\)") 
    cllmatch = re.search(CLLRegex, text)
    cll = cllmatch.group() if cllmatch else "No match" 

    priceRegex = r"INVOICE AMOUNT\s*:\s*EUR\s*([0-9,.]+)"
    pricematch = re.search(priceRegex, text)
    price = pricematch.group(1) if pricematch else "No match"



    vesselRegex = r" CARRIER \s*:\s*([A-Z]*\s*[A-Z]*\s*[0-9A-Z]+)"
    vesselmatch = re.search(vesselRegex, text)
    vessel = vesselmatch.group(1) if vesselmatch else "No match"

    

    CTRREGEX = re.compile(r"(?!" + re.escape(bl[1:11]) + r")[A-Z][A-Z][A-Z][A-Z][1-9][0-9][0-9][0-9][0-9][0-9][0-9]")
    ctr_matches = re.findall(CTRREGEX, text)
    ctr = ctr_matches[0] if ctr_matches else "No match"
 
    insuranceRegex = re.compile("[A-Z$][0-9][0-9][0-9][0-9][A-Z][A-Z0-9][0-9][0-9][0-9][0-9][0-9]")
    insurancematch = re.search(insuranceRegex,text)
    insurance = insurancematch.group() if insurancematch else "No match"

    #0.4 van hogy az S-t $-nek olvassa es az I-t 1-nek. 
    if insurance.startswith("$"):
        insurance = "S" + insurance[1:]
    if insurance[6] == "1":
        insurance = insurance[0:6] + 'I' + insurance[7:12]



    replace_dict = {
        "{{szamla}}": szamla,
        "{{price}": price,
        "{{bl}}":  bl,
	"{{cll}}":  cll,
        "{{today}}": str(now), 
        "{{ctr}}" : ctr,    #0.3
        "{{insurance}}" : insurance
    }

    wb = openpyxl.load_workbook('excel.xlsx')
    sheet = wb.active
    sheet.title = 'MyPDF'
    
    row = sheet.max_row + 1
    sheet.cell(row=row, column=2).value = ctr
    sheet.cell(row=row, column=17).value = "Korea"
    sheet.cell(row=row, column=18).value = vessel
    sheet.cell(row=row, column=21).value = szamla
    sheet.cell(row=row, column=23).value = bl
    sheet.cell(row=row, column=25).value = price
    sheet.cell(row=row, column=26).value = cll
    sheet.cell(row=row, column=27).value = insurance
    wb.save('excel.xlsx')

    docx = Document("SOD_.docx")
    docx = replace_text_with_variables(docx, replace_dict)
    if szamla == "No match":
        docx.save(f"{current_directory}/{folder_name}/SOD_{szamla}_No match szamla.docx")
    if price == "No match":
        docx.save(f"{current_directory}/{folder_name}/SOD_{szamla}_No match price.docx")
    if bl == "No match":
        docx.save(f"{current_directory}/{folder_name}/SOD_{szamla}_No match bl.docx")
    if cll == "No match":
        docx.save(f"{current_directory}/{folder_name}/SOD_{szamla}_No match cll.docx")
    if ctr == "No match":
        docx.save(f"{current_directory}/{folder_name}/SOD_{szamla}_No match ctr.docx")
    else:
        docx.save(f"{current_directory}/{folder_name}/SOD_{szamla}.docx")

    
    docx2 = Document("IMPORT ORDER_.docx")
    docx2 = replace_text_with_variables(docx2, replace_dict)
    if price == "No match":
        docx2.save(f"{current_directory}/{folder_name}/IMPORT ORDER_{szamla}_No match price.docx")
    if insurance == "No match":
        docx2.save(f"{current_directory}/{folder_name}/IMPORT ORDER_{szamla}_No match insurance.docx")
    else:
        docx2.save(f"{current_directory}/{folder_name}/IMPORT ORDER_{szamla}.docx")

    try:
        os.rename(full_pdf_path, os.path.join(full_path, pdf_file))
    except Exception as e:
        print(e)